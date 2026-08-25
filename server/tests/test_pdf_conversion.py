from io import BytesIO
from pathlib import Path
import subprocess
import sys

import pytest
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader, PdfWriter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Table, TableStyle

from app.errors import AppError
from app.services.pdf_conversion import (
    PAYMENT_VARIANCE_COLUMN_STARTS,
    PAYMENT_VARIANCE_HEADERS,
    PdfConversionRunner,
    PdfConversionService,
)


def _write_text_pdf(path: Path, page_count: int = 1) -> None:
    document = canvas.Canvas(str(path), pagesize=A4)
    for page_number in range(page_count):
        document.drawString(72, 760, f"Editable page {page_number + 1}")
        document.showPage()
    document.save()


def _write_blank_pdf(path: Path) -> None:
    document = canvas.Canvas(str(path), pagesize=A4)
    document.rect(72, 600, 100, 100)
    document.save()


def _write_table_pdf(path: Path) -> None:
    document = canvas.Canvas(str(path), pagesize=A4)
    document.drawString(72, 760, "Quarterly summary")
    document.drawString(72, 730, "Amounts are shown in RMB.")
    table = Table(
        [["Name", "Amount"], ["Alpha", "12.50"], ["Beta", "-3.25"]],
        colWidths=[140, 100],
        rowHeights=[24, 24, 24],
    )
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.8, colors.black),
                ("FONT", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        )
    )
    table.wrapOn(document, 300, 200)
    table.drawOn(document, 72, 620)
    document.save()


def _write_payment_variance_pdf(path: Path) -> None:
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    except KeyError:
        pass
    page_width, page_height = landscape(A4)
    document = canvas.Canvas(str(path), pagesize=(page_width, page_height))
    document.setTitle("Sa_Rpt121_TW.rpt")
    document.setFont("STSong-Light", 5.5)
    document.drawString(
        10,
        page_height - 18,
        "付款過數日期: 01/07/2026--31/07/2026",
    )
    document.drawString(520, page_height - 18, "程式編號: SA_RPT121")
    document.drawString(
        10,
        page_height - 28,
        "只顯示差異數據: 是 測 試 有 限 公 司 用戶: ADMIN",
    )
    document.drawString(520, page_height - 28, "日期: 17/07/2026")
    document.drawString(520, page_height - 38, "時間: 09:31:02")

    def draw_headers():
        for x, heading in zip(PAYMENT_VARIANCE_COLUMN_STARTS, PAYMENT_VARIANCE_HEADERS):
            document.drawString(x, page_height - 84, heading)

    def draw_values(y, values):
        for index, value in values.items():
            document.drawString(PAYMENT_VARIANCE_COLUMN_STARTS[index], y, str(value))

    draw_headers()
    draw_values(
        page_height - 108,
        {
            0: "9010005",
            1: "Acme Packaging",
            2: "Manager A",
            3: "Q001",
            4: "25/05/2026",
            5: "31",
            6: "MONTHLY 25",
            7: "INV001",
            8: "27/05/2026",
            9: "07/07/2026",
            10: "RMB",
            11: "01/07/2026",
            12: "DRAFT001",
            13: "29/12/2026",
            14: "175",
            15: "電匯",
            16: "6個月承兌匯票",
            17: "24.323",
            18: "75,887.76",
            19: "24.323",
            20: "75,887.76",
            21: "729.69",
            22: "0.00",
            23: "未補差價",
            24: "等待核准",
        },
    )
    draw_values(page_height - 114, {1: "Ltd", 12: "PART2"})
    draw_values(
        page_height - 132,
        {16: "9010005_客戶小計:", 17: "24.323", 18: "75,887.76", 21: "729.69"},
    )
    document.showPage()

    document.setFont("STSong-Light", 5.5)
    draw_headers()
    draw_values(
        page_height - 108,
        {
            0: "9010006",
            1: "Second Customer",
            2: "Manager B",
            3: "Q002",
            4: "26/05/2026",
            5: "30",
            6: "MONTHLY 30",
            7: "INV002",
            8: "28/05/2026",
            9: "08/07/2026",
            10: "RMB",
            17: "10.000",
            18: "20,000.00",
            19: "9.500",
            20: "19,000.00",
            21: "-100.00",
            22: "0.00",
        },
    )
    draw_values(
        page_height - 132,
        {
            16: "總計:",
            17: "34.323",
            18: "95,887.76",
            19: "33.823",
            20: "94,887.76",
            21: "629.69",
            22: "0.00",
        },
    )
    document.save()


def test_generic_table_converts_to_xlsx(work_dir):
    source = work_dir / "table.pdf"
    output = work_dir / "table.xlsx"
    _write_table_pdf(source)

    PdfConversionService(max_pages=30).convert(source, output, "xlsx")

    sheet = load_workbook(output).active
    assert sheet["A1"].value == "Name"
    assert sheet["B2"].value == 12.5
    assert sheet.freeze_panes == "A2"
    assert sheet.auto_filter.ref


def test_generic_table_converts_to_docx_with_surrounding_text(work_dir):
    source = work_dir / "table.pdf"
    output = work_dir / "table.docx"
    _write_table_pdf(source)

    PdfConversionService(max_pages=30).convert(source, output, "docx")

    document = Document(output)
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Quarterly summary" in paragraphs
    assert "Amounts are shown in RMB." in paragraphs
    assert document.tables[0].cell(1, 0).text == "Alpha"


def test_payment_variance_report_merges_pages_and_continuations(work_dir):
    source = work_dir / "variance.pdf"
    output = work_dir / "variance.xlsx"
    _write_payment_variance_pdf(source)

    PdfConversionService(max_pages=30).convert(source, output, "xlsx")

    sheet = load_workbook(output).active
    assert sheet.max_column == 25
    assert sum(
        1
        for row in sheet.iter_rows(values_only=True)
        if row and row[0] == PAYMENT_VARIANCE_HEADERS[0]
    ) == 1
    assert sheet["A1"].value == "測試有限公司"
    assert sheet["A2"].value == "付款類型差異報表"
    assert "01/07/2026--31/07/2026" in sheet["A3"].value
    assert sheet["A6"].value == PAYMENT_VARIANCE_HEADERS[0]
    assert sheet["A7"].value == "9010005"
    assert "Ltd" in sheet["B7"].value
    assert "PART2" in sheet["M7"].value
    assert sheet["E7"].value.year == 2026
    assert sheet["R7"].value == 24.323
    assert sheet["S7"].value == 75887.76
    assert any(
        "總計" in "".join(str(cell.value or "") for cell in row)
        for row in sheet.iter_rows()
    )


def test_payment_variance_report_converts_to_editable_docx(work_dir):
    source = work_dir / "variance.pdf"
    output = work_dir / "variance.docx"
    _write_payment_variance_pdf(source)

    PdfConversionService(max_pages=30).convert(source, output, "docx")

    document = Document(output)
    assert document.tables
    assert len(document.tables[0].columns) == 25
    assert document.tables[0].cell(1, 0).text == "9010005"


def test_rejects_scanned_pdf(work_dir):
    source = work_dir / "blank.pdf"
    _write_blank_pdf(source)

    with pytest.raises(AppError) as error:
        PdfConversionService(max_pages=30).convert(
            source, work_dir / "blank.docx", "docx"
        )

    assert error.value.code == "SCANNED_PDF_UNSUPPORTED"


def test_rejects_encrypted_pdf(work_dir):
    plain = work_dir / "plain.pdf"
    encrypted = work_dir / "encrypted.pdf"
    _write_text_pdf(plain)
    reader = PdfReader(plain)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt("secret")
    with encrypted.open("wb") as stream:
        writer.write(stream)

    with pytest.raises(AppError) as error:
        PdfConversionService(max_pages=30).convert(
            encrypted, work_dir / "encrypted.docx", "docx"
        )

    assert error.value.code == "ENCRYPTED_PDF"


def test_rejects_too_many_pages(work_dir):
    source = work_dir / "pages.pdf"
    _write_text_pdf(source, page_count=2)

    with pytest.raises(AppError) as error:
        PdfConversionService(max_pages=1).convert(
            source, work_dir / "pages.docx", "docx"
        )

    assert error.value.code == "PDF_TOO_MANY_PAGES"


def test_runner_maps_timeout(monkeypatch, work_dir):
    runner = PdfConversionRunner(max_pages=30, timeout_seconds=0.01)
    monkeypatch.setattr(
        PdfConversionRunner,
        "ready",
        property(lambda self: True),
    )

    def timeout(*args, **kwargs):
        raise subprocess.TimeoutExpired(args[0], 0.01)

    monkeypatch.setattr(subprocess, "run", timeout)
    with pytest.raises(AppError) as error:
        runner.run(work_dir / "in.pdf", work_dir / "out.xlsx", "xlsx")
    assert error.value.code == "CONVERSION_TIMEOUT"


def test_runner_uses_in_process_conversion_when_frozen(monkeypatch, work_dir):
    runner = PdfConversionRunner(max_pages=30, timeout_seconds=1)
    monkeypatch.setattr(sys, "frozen", True, raising=False)
    monkeypatch.setattr(
        PdfConversionRunner,
        "ready",
        property(lambda self: True),
    )
    converted = []

    def convert(self, input_path, output_path, output_format):
        converted.append((input_path, output_path, output_format))

    monkeypatch.setattr(PdfConversionService, "convert", convert)
    source = work_dir / "in.pdf"
    output = work_dir / "out.xlsx"
    runner.run(source, output, "xlsx")
    assert converted == [(source, output, "xlsx")]
