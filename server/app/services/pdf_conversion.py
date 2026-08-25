from __future__ import annotations

import bisect
import importlib.util
import json
import re
import subprocess
import sys
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Literal

from ..errors import AppError

OutputFormat = Literal["xlsx", "docx"]

PAYMENT_VARIANCE_HEADERS = [
    "客戶編號",
    "客戶名稱",
    "營銷經理",
    "報價單",
    "報價單日期",
    "截數日",
    "付款方式",
    "發票編號",
    "發票日期",
    "發票到期日",
    "貨幣",
    "收款日期",
    "承兌匯票編號",
    "匯票到期日",
    "承兌天數",
    "合同收款類型",
    "實際收款類型",
    "重量(噸)",
    "發票金額",
    "收款重量(噸)",
    "收款金額",
    "差異金額",
    "已補金額",
    "差異回收情況",
    "狀態",
]

PAYMENT_VARIANCE_COLUMN_STARTS = [
    11.6,
    35.6,
    104.6,
    128.6,
    161.6,
    192.4,
    209.8,
    241.6,
    281.4,
    313.6,
    345.9,
    368.9,
    392.5,
    431.7,
    463.2,
    486.4,
    521.1,
    565.0,
    606.4,
    633.6,
    686.6,
    722.4,
    753.5,
    777.5,
    813.5,
]

PAYMENT_VARIANCE_COLUMN_BOUNDS = [
    0.0,
    34.0,
    102.0,
    127.0,
    160.0,
    191.0,
    208.5,
    240.0,
    280.5,
    312.0,
    344.0,
    360.0,
    391.0,
    430.0,
    461.0,
    485.0,
    520.0,
    560.0,
    594.0,
    632.0,
    670.0,
    716.0,
    752.0,
    781.0,
    811.0,
]

PAYMENT_VARIANCE_COLUMN_WIDTHS = [
    11,
    28,
    12,
    15,
    13,
    8,
    16,
    22,
    13,
    13,
    8,
    13,
    24,
    13,
    10,
    16,
    18,
    12,
    15,
    14,
    15,
    14,
    13,
    18,
    14,
]


@dataclass
class LogicalRow:
    values: list[str]
    kind: Literal["detail", "summary"]


@dataclass(frozen=True)
class PaymentVarianceMetadata:
    company_name: str = ""
    report_title: str = "付款類型差異報表"
    payment_period: str = ""
    difference_only: str = ""
    program_number: str = ""
    user_name: str = ""
    report_date: str = ""
    report_time: str = ""


class PdfConversionService:
    def __init__(self, max_pages: int) -> None:
        self.max_pages = max_pages

    def convert(self, input_path: Path, output_path: Path, output_format: OutputFormat) -> None:
        self._validate_pdf(input_path)
        if output_format == "xlsx":
            self._convert_to_xlsx(input_path, output_path)
        elif output_format == "docx":
            self._convert_to_docx(input_path, output_path)
        else:
            raise AppError("INVALID_REQUEST", "輸出格式必須是 xlsx 或 docx", 422)

    def _validate_pdf(self, input_path: Path) -> None:
        try:
            from pypdf import PdfReader
            from pypdf.errors import PdfReadError

            if input_path.read_bytes()[:5] != b"%PDF-":
                raise AppError("INVALID_PDF", "文件不是有效的 PDF", 400)
            reader = PdfReader(str(input_path), strict=False)
            if reader.is_encrypted:
                raise AppError("ENCRYPTED_PDF", "暫不支持加密或需要密碼的 PDF", 422)
            page_count = len(reader.pages)
            if page_count == 0:
                raise AppError("INVALID_PDF", "PDF 沒有可轉換的頁面", 400)
            if page_count > self.max_pages:
                raise AppError(
                    "PDF_TOO_MANY_PAGES",
                    f"PDF 不能超過 {self.max_pages} 頁",
                    413,
                )
        except AppError:
            raise
        except (PdfReadError, OSError, ValueError, TypeError) as exc:
            raise AppError("INVALID_PDF", "PDF 已損壞或格式不正確", 400) from exc

    def _open_pdf(self, input_path: Path):
        try:
            import pdfplumber

            pdf = pdfplumber.open(str(input_path))
            if not any((page.extract_text() or "").strip() for page in pdf.pages):
                pdf.close()
                raise AppError(
                    "SCANNED_PDF_UNSUPPORTED",
                    "PDF 沒有可選文字，暫不支持掃描件 OCR",
                    422,
                )
            return pdf
        except AppError:
            raise
        except Exception as exc:
            raise AppError("INVALID_PDF", "無法讀取 PDF 內容", 400) from exc

    def _convert_to_xlsx(self, input_path: Path, output_path: Path) -> None:
        pdf = self._open_pdf(input_path)
        try:
            if self._is_payment_variance_report(pdf):
                rows = self._extract_payment_variance_rows(pdf)
                if not rows:
                    raise AppError("NO_TABLE_FOUND", "PDF 中沒有可轉換的表格", 422)
                metadata = self._extract_payment_variance_metadata(pdf)
                self._write_payment_variance_workbook(rows, metadata, output_path)
                return

            tables = self._extract_generic_tables(pdf)
            if not tables:
                raise AppError(
                    "NO_TABLE_FOUND",
                    "PDF 中沒有識別到表格，請改為轉 Word",
                    422,
                )
            self._write_generic_workbook(tables, output_path)
        finally:
            pdf.close()

    def _convert_to_docx(self, input_path: Path, output_path: Path) -> None:
        pdf = self._open_pdf(input_path)
        try:
            if self._is_payment_variance_report(pdf):
                rows = self._extract_payment_variance_rows(pdf)
                self._write_payment_variance_document(rows, output_path)
            else:
                self._write_generic_document(pdf, output_path)
        finally:
            pdf.close()

    @staticmethod
    def _is_payment_variance_report(pdf) -> bool:
        metadata_title = str((pdf.metadata or {}).get("Title", "")).lower()
        first_text = (pdf.pages[0].extract_text() or "").replace(" ", "")
        return "sa_rpt121" in metadata_title or any(
            title in first_text for title in ("付款類型差異報表", "付款类型差异报表")
        )

    def _extract_payment_variance_rows(self, pdf) -> list[LogicalRow]:
        extracted: list[LogicalRow] = []
        for page in pdf.pages:
            scale = float(page.width) / 841.95
            bounds = [value * scale for value in PAYMENT_VARIANCE_COLUMN_BOUNDS]
            data_top = float(page.height) * 0.143
            words = [
                word
                for word in page.extract_words(
                    x_tolerance=0.1,
                    y_tolerance=1,
                    keep_blank_chars=False,
                )
                if data_top <= float(word["top"]) <= float(page.height) - 8
            ]
            physical_lines = self._cluster_items_by_line(words)
            current: LogicalRow | None = None

            for line_words in physical_lines:
                cells = ["" for _ in bounds]
                for word in sorted(line_words, key=lambda item: float(item["x0"])):
                    text = str(word.get("text", "")).strip()
                    center = (float(word["x0"]) + float(word["x1"])) / 2
                    if re.fullmatch(r"-?[\d,.]+", text) and 752 * scale <= center < 811 * scale:
                        column = 22
                    elif not re.search(r"\d", text) and float(word["x0"]) >= 811 * scale:
                        column = 24
                    elif not re.search(r"\d", text) and float(word["x0"]) >= 777 * scale:
                        column = 23
                    else:
                        column = bisect.bisect_right(bounds, center) - 1
                    if 0 <= column < len(cells):
                        separator = " " if cells[column] else ""
                        cells[column] += f"{separator}{text}"
                cells = [self._clean_cell_text(value) for value in cells]
                combined = "".join(cells)
                if not combined or "報表結束" in combined or "报表结束" in combined:
                    continue

                first_cell = re.sub(r"\s+", "", cells[0])
                is_detail = bool(re.fullmatch(r"\d{7}", first_cell))
                is_summary = bool(re.search(r"(小計|小计|總計|总计)", combined))

                if is_detail or is_summary:
                    if is_summary:
                        self._normalize_summary_label(cells)
                    if current is not None:
                        extracted.append(current)
                    current = LogicalRow(
                        values=cells,
                        kind="detail" if is_detail else "summary",
                    )
                    continue

                if current is None:
                    continue
                if current.kind == "summary" and not re.search(r"\d", combined):
                    continue
                self._append_continuation(current.values, cells)

            if current is not None:
                extracted.append(current)
        return extracted

    @staticmethod
    def _extract_payment_variance_metadata(pdf) -> PaymentVarianceMetadata:
        text = pdf.pages[0].extract_text() or ""

        def match(pattern: str, flags: int = 0) -> str:
            result = re.search(pattern, text, flags)
            return result.group(1).strip() if result else ""

        company_name = match(
            r"(?:是|否)\s+(.+?)\s+(?:用戶|用户)\s*[:：]",
            re.S,
        )
        company_name = re.sub(
            r"(?<=[\u3400-\u9fff])\s+(?=[\u3400-\u9fff])",
            "",
            company_name,
        )
        company_name = re.sub(r"\s{2,}", " ", company_name).strip()
        report_title = (
            "付款类型差异报表"
            if "付款类型差异报表" in text
            else "付款類型差異報表"
        )
        return PaymentVarianceMetadata(
            company_name=company_name,
            report_title=report_title,
            payment_period=match(
                r"付款(?:過數|过数)日期\s*[:：]\s*([^\s]+)"
            ),
            difference_only=match(
                r"只(?:顯示|显示)差異數據\s*[:：]\s*([^\s]+)"
            ),
            program_number=match(r"程式編號\s*[:：]\s*([^\s]+)"),
            user_name=match(r"(?:用戶|用户)\s*[:：]\s*([^\s]+)"),
            report_date=match(r"^日期\s*[:：]\s*([^\s]+)", re.M),
            report_time=match(r"^時間\s*[:：]\s*([^\s]+)", re.M),
        )

    @staticmethod
    def _normalize_summary_label(cells: list[str]) -> None:
        label = "".join(value for value in cells[:17] if value)
        if not label:
            return
        for index in range(17):
            cells[index] = ""
        cells[16] = label

    @staticmethod
    def _cluster_items_by_line(items: list[dict]) -> list[list[dict]]:
        lines: list[list[dict]] = []
        centers: list[float] = []
        for item in sorted(items, key=lambda value: (float(value["top"]), float(value["x0"]))):
            top = float(item["top"])
            if not lines or top - centers[-1] > 3.0:
                lines.append([item])
                centers.append(top)
            else:
                lines[-1].append(item)
                centers[-1] = sum(float(item["top"]) for item in lines[-1]) / len(lines[-1])
        return lines

    @staticmethod
    def _clean_cell_text(value: str) -> str:
        return re.sub(r"[ \t]+", " ", value).strip()

    @staticmethod
    def _append_continuation(target: list[str], continuation: list[str]) -> None:
        for index, value in enumerate(continuation):
            if not value:
                continue
            if not target[index]:
                target[index] = value
            elif value not in target[index].splitlines():
                target[index] = f"{target[index]}\n{value}"

    def _extract_generic_tables(self, pdf) -> list[tuple[str, list[list[str]]]]:
        merged: dict[tuple[str, ...], tuple[str, list[list[str]]]] = {}
        sequence = 1
        for page in pdf.pages:
            for table_number, rows in enumerate(self._find_page_tables(page), start=1):
                header_key = tuple(self._normalize_header(value) for value in rows[0])
                if header_key in merged:
                    merged[header_key][1].extend(rows[1:])
                    continue
                name = f"P{page.page_number}_T{table_number or sequence}"
                merged[header_key] = (name, [list(row) for row in rows])
                sequence += 1
        return list(merged.values())

    def _find_page_tables(self, page) -> list[list[list[str]]]:
        return [rows for rows, _ in self._find_page_table_regions(page)]

    def _find_page_table_regions(
        self, page
    ) -> list[tuple[list[list[str]], tuple[float, float, float, float]]]:
        settings_candidates = [
            {
                "vertical_strategy": "lines",
                "horizontal_strategy": "lines",
                "snap_tolerance": 3,
                "join_tolerance": 3,
            },
            {
                "vertical_strategy": "text",
                "horizontal_strategy": "text",
                "min_words_vertical": 3,
                "min_words_horizontal": 2,
                "intersection_tolerance": 4,
            },
        ]
        for settings in settings_candidates:
            tables: list[
                tuple[list[list[str]], tuple[float, float, float, float]]
            ] = []
            try:
                found_tables = page.find_tables(settings)
            except Exception:
                continue
            for table in found_tables:
                cleaned = self._clean_table(table.extract())
                if len(cleaned) >= 2 and len(cleaned[0]) >= 2:
                    tables.append(
                        (
                            cleaned,
                            tuple(float(value) for value in table.bbox),
                        )
                    )
            if tables:
                return tables
        return []

    def _clean_table(self, rows) -> list[list[str]]:
        normalized = [
            [self._clean_cell_text(str(value or "")) for value in row]
            for row in rows
            if any(str(value or "").strip() for value in row)
        ]
        if not normalized:
            return []
        width = max(len(row) for row in normalized)
        padded = [row + [""] * (width - len(row)) for row in normalized]
        keep_columns = [
            index
            for index in range(width)
            if any(row[index].strip() for row in padded)
        ]
        return [[row[index] for index in keep_columns] for row in padded]

    @staticmethod
    def _normalize_header(value: str) -> str:
        return re.sub(r"\s+", "", value or "").lower()

    def _write_payment_variance_workbook(
        self,
        rows: list[LogicalRow],
        metadata: PaymentVarianceMetadata,
        output_path: Path,
    ) -> None:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font

        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "付款差異明細"
        worksheet.append([metadata.company_name])
        worksheet.append([metadata.report_title])
        worksheet.append(
            [
                f"付款過數日期：{metadata.payment_period}"
                if metadata.payment_period
                else "",
                *[""] * 15,
                f"程式編號：{metadata.program_number}"
                if metadata.program_number
                else "",
            ]
        )
        left_detail = (
            f"只顯示差異數據：{metadata.difference_only}"
            if metadata.difference_only
            else ""
        )
        right_detail = "  ".join(
            value
            for value in (
                f"用戶：{metadata.user_name}" if metadata.user_name else "",
                f"日期：{metadata.report_date}" if metadata.report_date else "",
                f"時間：{metadata.report_time}" if metadata.report_time else "",
            )
            if value
        )
        worksheet.append([left_detail, *([""] * 15), right_detail])
        worksheet.append([])
        worksheet.append(PAYMENT_VARIANCE_HEADERS)
        worksheet.merge_cells("A1:Y1")
        worksheet.merge_cells("A2:Y2")
        worksheet.merge_cells("A3:P3")
        worksheet.merge_cells("Q3:Y3")
        worksheet.merge_cells("A4:P4")
        worksheet.merge_cells("Q4:Y4")
        worksheet["A1"].font = Font(
            name="Microsoft YaHei", size=18, bold=True, color="1F2937"
        )
        worksheet["A1"].alignment = Alignment(horizontal="center", vertical="center")
        worksheet["A2"].font = Font(
            name="Microsoft YaHei", size=15, bold=True, color="1F2937"
        )
        worksheet["A2"].alignment = Alignment(horizontal="center", vertical="center")
        for cell_reference in ("A3", "Q3", "A4", "Q4"):
            worksheet[cell_reference].font = Font(
                name="Microsoft YaHei", size=9, color="334155"
            )
            worksheet[cell_reference].alignment = Alignment(vertical="center")
        worksheet["Q3"].alignment = Alignment(horizontal="right", vertical="center")
        worksheet["Q4"].alignment = Alignment(horizontal="right", vertical="center")
        worksheet.row_dimensions[1].height = 28
        worksheet.row_dimensions[2].height = 24
        worksheet.row_dimensions[3].height = 18
        worksheet.row_dimensions[4].height = 18
        worksheet.row_dimensions[5].height = 6
        row_kinds: list[str] = []
        for logical_row in rows:
            worksheet.append(
                [
                    self._payment_variance_value(index, value)
                    for index, value in enumerate(logical_row.values)
                ]
            )
            row_kinds.append(logical_row.kind)
        self._style_worksheet(
            worksheet,
            row_kinds=row_kinds,
            column_widths=PAYMENT_VARIANCE_COLUMN_WIDTHS,
            difference_column=22,
            header_row=6,
        )
        workbook.save(output_path)

    def _write_generic_workbook(
        self, tables: list[tuple[str, list[list[str]]]], output_path: Path
    ) -> None:
        from openpyxl import Workbook

        workbook = Workbook()
        workbook.remove(workbook.active)
        used_names: set[str] = set()
        for suggested_name, rows in tables:
            title = self._unique_sheet_name(suggested_name, used_names)
            worksheet = workbook.create_sheet(title)
            for row_index, row in enumerate(rows):
                headers = rows[0] if rows else []
                worksheet.append(
                    [
                        self._generic_value(value, headers[column] if column < len(headers) else "")
                        for column, value in enumerate(row)
                    ]
                )
            widths = [
                min(36, max(10, max(len(str(row[col] or "")) for row in rows) + 2))
                for col in range(len(rows[0]))
            ]
            self._style_worksheet(worksheet, column_widths=widths)
        workbook.save(output_path)

    @staticmethod
    def _style_worksheet(
        worksheet,
        row_kinds: list[str] | None = None,
        column_widths: list[float] | None = None,
        difference_column: int | None = None,
        header_row: int = 1,
    ) -> None:
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter

        header_fill = PatternFill("solid", fgColor="D9EAF7")
        summary_fill = PatternFill("solid", fgColor="D9EAF7")
        total_fill = PatternFill("solid", fgColor="B4C6E7")
        header_border = Border(bottom=Side(style="medium", color="17365D"))
        row_border = Border(bottom=Side(style="hair", color="D9E2F3"))

        worksheet.sheet_view.showGridLines = False
        worksheet.sheet_view.zoomScale = 70
        data_start_row = header_row + 1
        worksheet.freeze_panes = f"A{data_start_row}"
        last_column = get_column_letter(worksheet.max_column)
        worksheet.auto_filter.ref = (
            f"A{header_row}:{last_column}{worksheet.max_row}"
        )
        worksheet.page_setup.orientation = "landscape"
        worksheet.page_setup.fitToWidth = 1
        worksheet.sheet_properties.pageSetUpPr.fitToPage = True

        for cell in worksheet[header_row]:
            cell.fill = header_fill
            cell.font = Font(name="Microsoft YaHei", size=9, bold=True, color="1F2937")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = header_border
        worksheet.row_dimensions[header_row].height = 32

        for row_index, row in enumerate(
            worksheet.iter_rows(min_row=data_start_row),
            start=data_start_row,
        ):
            kind = (
                row_kinds[row_index - data_start_row]
                if row_kinds
                else "detail"
            )
            row_text = "".join(str(cell.value or "") for cell in row)
            is_total = kind == "summary" and bool(re.search(r"(總計|总计)", row_text))
            max_lines = 1
            for cell in row:
                max_lines = max(max_lines, str(cell.value or "").count("\n") + 1)
                cell.font = Font(
                    name="Microsoft YaHei",
                    size=9,
                    bold=kind == "summary",
                )
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = row_border
                if isinstance(cell.value, (date, datetime)):
                    cell.number_format = "dd/mm/yyyy"
                elif isinstance(cell.value, float):
                    cell.number_format = (
                        "#,##0.000"
                        if cell.column in {18, 20}
                        else "#,##0.00"
                    )
                elif isinstance(cell.value, int) and cell.column not in {1, 4, 8, 13}:
                    cell.number_format = "#,##0"
                if kind == "summary":
                    cell.fill = total_fill if is_total else summary_fill
            worksheet.row_dimensions[row_index].height = min(72, 18 * max_lines)
            if difference_column is not None:
                value = worksheet.cell(row=row_index, column=difference_column).value
                if isinstance(value, (int, float)) and value < 0:
                    worksheet.cell(row=row_index, column=difference_column).font = Font(
                        name="Microsoft YaHei",
                        size=9,
                        bold=kind == "summary",
                        color="C00000",
                    )

        if column_widths:
            for index, width in enumerate(column_widths, start=1):
                worksheet.column_dimensions[get_column_letter(index)].width = width

    def _payment_variance_value(self, index: int, value: str):
        if not value:
            return None
        if "\n" in value:
            return value
        if index in {4, 8, 9, 11, 13}:
            parsed = self._parse_date(value)
            return parsed or value
        if index in {5, 14, 17, 18, 19, 20, 21, 22}:
            parsed = self._parse_number(value)
            return parsed if parsed is not None else value
        return value

    def _generic_value(self, value: str, header: str):
        if not value or "\n" in value:
            return value or None
        if re.search(r"(編號|编号|代碼|代码|單號|单号|帳號|账号|ID)", header, re.I):
            return value
        parsed_date = self._parse_date(value)
        if parsed_date:
            return parsed_date
        parsed_number = self._parse_number(value)
        return parsed_number if parsed_number is not None else value

    @staticmethod
    def _parse_date(value: str) -> date | None:
        for pattern in ("%d/%m/%Y", "%Y-%m-%d", "%Y/%m/%d"):
            try:
                return datetime.strptime(value.strip(), pattern).date()
            except ValueError:
                continue
        return None

    @staticmethod
    def _parse_number(value: str) -> int | float | None:
        cleaned = value.strip().replace(",", "")
        if not re.fullmatch(r"-?(?:\d+|\d*\.\d+)", cleaned):
            return None
        if cleaned.startswith("0") and len(cleaned) > 1 and not cleaned.startswith("0."):
            return None
        number = float(cleaned) if "." in cleaned else int(cleaned)
        return number

    @staticmethod
    def _unique_sheet_name(suggested: str, used: set[str]) -> str:
        base = re.sub(r"[\[\]:*?/\\]", "_", suggested or "Table")[:31] or "Table"
        candidate = base
        counter = 2
        while candidate in used:
            suffix = f"_{counter}"
            candidate = f"{base[:31-len(suffix)]}{suffix}"
            counter += 1
        used.add(candidate)
        return candidate

    def _write_payment_variance_document(
        self, rows: list[LogicalRow], output_path: Path
    ) -> None:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt

        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
        section.top_margin = Mm(7)
        section.bottom_margin = Mm(7)
        section.left_margin = Mm(6)
        section.right_margin = Mm(6)

        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(6)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(5)
        run = title.add_run("付款類型差異報表")
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(14)

        table = document.add_table(rows=1, cols=len(PAYMENT_VARIANCE_HEADERS))
        table.autofit = False
        table.style = "Table Grid"
        usable_width_dxa = int((297 - 12) / 25.4 * 1440)
        source_widths = [
            PAYMENT_VARIANCE_COLUMN_STARTS[index + 1]
            - PAYMENT_VARIANCE_COLUMN_STARTS[index]
            if index + 1 < len(PAYMENT_VARIANCE_COLUMN_STARTS)
            else 841.95 - PAYMENT_VARIANCE_COLUMN_STARTS[index]
            for index in range(len(PAYMENT_VARIANCE_COLUMN_STARTS))
        ]
        total_source_width = sum(source_widths)
        column_widths = [
            max(250, int(usable_width_dxa * width / total_source_width))
            for width in source_widths
        ]
        self._set_table_width(table, sum(column_widths), column_widths)

        header_cells = table.rows[0].cells
        for index, heading in enumerate(PAYMENT_VARIANCE_HEADERS):
            self._set_docx_cell(
                header_cells[index],
                heading,
                font_size=5.5,
                bold=True,
                fill="D9EAF7",
            )
        self._repeat_table_header(table.rows[0])

        for logical_row in rows:
            row = table.add_row()
            combined = "".join(logical_row.values)
            fill = "B4C6E7" if re.search(r"(總計|总计)", combined) else "EAF2F8"
            for index, value in enumerate(logical_row.values):
                self._set_docx_cell(
                    row.cells[index],
                    value,
                    font_size=5,
                    bold=logical_row.kind == "summary",
                    fill=fill if logical_row.kind == "summary" else None,
                )
        document.save(output_path)

    def _write_generic_document(self, pdf, output_path: Path) -> None:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.shared import Inches, Pt

        document = Document()
        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10)
        normal.paragraph_format.space_after = Pt(4)
        normal.paragraph_format.line_spacing = 1.15

        for page_index, page in enumerate(pdf.pages):
            section = (
                document.sections[0]
                if page_index == 0
                else document.add_section(WD_SECTION.NEW_PAGE)
            )
            section.page_width = Pt(float(page.width))
            section.page_height = Pt(float(page.height))
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

            table_regions = self._find_page_table_regions(page)
            if table_regions:
                table_boxes = [bbox for _, bbox in table_regions]
                words = [
                    word
                    for word in page.extract_words(
                        x_tolerance=1,
                        y_tolerance=2,
                        keep_blank_chars=False,
                    )
                    if not any(
                        bbox[0]
                        <= (float(word["x0"]) + float(word["x1"])) / 2
                        <= bbox[2]
                        and bbox[1]
                        <= (float(word["top"]) + float(word["bottom"])) / 2
                        <= bbox[3]
                        for bbox in table_boxes
                    )
                ]
                elements: list[tuple[float, str, object]] = [
                    (
                        min(float(word["top"]) for word in line),
                        "text",
                        " ".join(
                            str(word["text"])
                            for word in sorted(
                                line, key=lambda value: float(value["x0"])
                            )
                        ),
                    )
                    for line in self._cluster_items_by_line(words)
                    if line
                ]
                elements.extend(
                    (bbox[1], "table", rows) for rows, bbox in table_regions
                )
                text_index = 0
                for _, element_type, content in sorted(
                    elements, key=lambda value: value[0]
                ):
                    if element_type == "table":
                        self._add_docx_table(document, content)
                        continue
                    line = str(content).strip()
                    if not line:
                        continue
                    paragraph = document.add_paragraph()
                    run = paragraph.add_run(line)
                    if text_index == 0:
                        run.bold = True
                        run.font.size = Pt(12)
                    text_index += 1
            else:
                for line in (page.extract_text(layout=True) or "").splitlines():
                    line = line.rstrip()
                    if line:
                        document.add_paragraph(line)
        document.save(output_path)

    def _add_docx_table(self, document, rows: list[list[str]]) -> None:
        if not rows or not rows[0]:
            return
        table = document.add_table(rows=0, cols=len(rows[0]))
        table.autofit = False
        table.style = "Table Grid"
        section = document.sections[-1]
        usable_width = int(section.page_width - section.left_margin - section.right_margin)
        widths = [max(400, usable_width // len(rows[0]))] * len(rows[0])
        self._set_table_width(table, sum(widths), widths)
        for row_index, values in enumerate(rows):
            row = table.add_row()
            for column, value in enumerate(values):
                self._set_docx_cell(
                    row.cells[column],
                    value,
                    font_size=7,
                    bold=row_index == 0,
                    fill="E8EEF5" if row_index == 0 else None,
                )
        if table.rows:
            self._repeat_table_header(table.rows[0])

    @staticmethod
    def _set_docx_cell(
        cell, text: str, font_size: float, bold: bool = False, fill: str | None = None
    ) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        cell.text = text or ""
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            for run in paragraph.runs:
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(font_size)
                run.bold = bold
        if fill:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), fill)
            cell._tc.get_or_add_tcPr().append(shading)
        margins = OxmlElement("w:tcMar")
        for edge, value in (("top", 40), ("start", 60), ("bottom", 40), ("end", 60)):
            margin = OxmlElement(f"w:{edge}")
            margin.set(qn("w:w"), str(value))
            margin.set(qn("w:type"), "dxa")
            margins.append(margin)
        cell._tc.get_or_add_tcPr().append(margins)

    @staticmethod
    def _set_table_width(table, total_width_dxa: int, column_widths: list[int]) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        table_properties = table._tbl.tblPr
        table_width = table_properties.first_child_found_in("w:tblW")
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            table_properties.append(table_width)
        table_width.set(qn("w:type"), "dxa")
        table_width.set(qn("w:w"), str(total_width_dxa))

        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        table_properties.append(layout)

        indent = OxmlElement("w:tblInd")
        indent.set(qn("w:type"), "dxa")
        indent.set(qn("w:w"), "60")
        table_properties.append(indent)

        for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, column_widths):
            grid_column.set(qn("w:w"), str(width))
        for column, width in zip(table.columns, column_widths):
            for cell in column.cells:
                tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tc_width.set(qn("w:type"), "dxa")
                tc_width.set(qn("w:w"), str(width))

    @staticmethod
    def _repeat_table_header(row) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        row_properties = row._tr.get_or_add_trPr()
        table_header = OxmlElement("w:tblHeader")
        table_header.set(qn("w:val"), "true")
        row_properties.append(table_header)


class PdfConversionRunner:
    _REQUIRED_MODULES = ("pypdf", "pdfplumber", "openpyxl", "docx")

    def __init__(self, max_pages: int, timeout_seconds: float) -> None:
        self.max_pages = max_pages
        self.timeout_seconds = timeout_seconds
        self.worker_path = Path(__file__).resolve().parents[1] / "conversion_worker.py"

    @property
    def ready(self) -> bool:
        worker_available = getattr(sys, "frozen", False) or self.worker_path.is_file()
        return worker_available and all(
            importlib.util.find_spec(module) is not None
            for module in self._REQUIRED_MODULES
        )

    def run(self, input_path: Path, output_path: Path, output_format: OutputFormat) -> None:
        if not self.ready:
            raise AppError("SERVER_BUSY", "PDF 轉換組件尚未就緒", 503)
        if getattr(sys, "frozen", False):
            PdfConversionService(self.max_pages).convert(
                input_path, output_path, output_format
            )
            return
        command = [
            sys.executable,
            str(self.worker_path),
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--format",
            output_format,
            "--max-pages",
            str(self.max_pages),
        ]
        creation_flags = subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
        try:
            completed = subprocess.run(
                command,
                check=False,
                capture_output=True,
                text=True,
                encoding="utf-8",
                timeout=self.timeout_seconds,
                creationflags=creation_flags,
            )
        except subprocess.TimeoutExpired as exc:
            raise AppError(
                "CONVERSION_TIMEOUT",
                "PDF 轉換超時，請減少頁數後重試",
                504,
            ) from exc
        if completed.returncode == 0 and output_path.is_file():
            return
        try:
            payload = json.loads(completed.stdout.strip() or "{}")
        except json.JSONDecodeError:
            payload = {}
        if completed.returncode == 2 and payload.get("code"):
            raise AppError(
                str(payload["code"]),
                str(payload.get("message") or "PDF 轉換失敗"),
                int(payload.get("statusCode") or 400),
            )
        raise AppError("INTERNAL_ERROR", "PDF 轉換失敗，請稍後重試", 500)
